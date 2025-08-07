"""
RAG Document Processing System
Implements document processing, chunking, and storage for RAG workflows
"""

from dataclasses import dataclass, field
from typing import List, Dict, Optional, Union, Any
from datetime import datetime
import uuid
import numpy as np
from pathlib import Path
import json
import sqlite3
from abc import ABC, abstractmethod

from app.logger import get_enhanced_logger
from .legal_protection import LegalRiskDetector, LegalChunkValidator, LegalRiskLevel
from .document_classifier import AdvancedDocumentClassifier, DocumentType
from .financial_chunker import FinancialDocumentChunker

logger = get_enhanced_logger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata and embeddings"""
    chunk_id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[np.ndarray] = None
    source_document_id: str = ""
    chunk_index: int = 0
    chunk_type: str = "text"  # text, table, image, code, etc.
    relevance_score: float = 0.0
    created_at: datetime = field(default_factory=datetime.utcnow)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for serialization"""
        return {
            'chunk_id': self.chunk_id,
            'content': self.content,
            'metadata': self.metadata,
            'source_document_id': self.source_document_id,
            'chunk_index': self.chunk_index,
            'chunk_type': self.chunk_type,
            'relevance_score': self.relevance_score,
            'created_at': self.created_at.isoformat(),
            'embedding_shape': self.embedding.shape if self.embedding is not None else None
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'DocumentChunk':
        """Create chunk from dictionary"""
        chunk = cls(
            chunk_id=data['chunk_id'],
            content=data['content'],
            metadata=data.get('metadata', {}),
            source_document_id=data['source_document_id'],
            chunk_index=data['chunk_index'],
            chunk_type=data.get('chunk_type', 'text'),
            relevance_score=data.get('relevance_score', 0.0),
            created_at=datetime.fromisoformat(data['created_at'])
        )
        return chunk


@dataclass
class Document:
    """Represents a document with metadata and chunks"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    filename: str = ""
    content: str = ""
    content_type: str = ""
    file_size: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    upload_date: datetime = field(default_factory=datetime.utcnow)
    processed_date: Optional[datetime] = None
    chunks: List[DocumentChunk] = field(default_factory=list)
    status: str = "pending"  # pending, processing, completed, error
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary for serialization"""
        return {
            'id': self.id,
            'filename': self.filename,
            'content_type': self.content_type,
            'file_size': self.file_size,
            'metadata': self.metadata,
            'upload_date': self.upload_date.isoformat(),
            'processed_date': self.processed_date.isoformat() if self.processed_date else None,
            'chunk_count': len(self.chunks),
            'status': self.status
        }


class DocumentProcessor:
    """Handles document processing for various file types"""
    
    def __init__(self):
        self.supported_types = {'.pdf', '.txt', '.docx', '.doc', '.html', '.md', '.json', '.xlsx', '.xls', '.csv'}
        self.logger = logger
    
    def process_document(self, content: Union[str, bytes], 
                        filename: str, 
                        content_type: str) -> Document:
        """
        Process a document and extract text content
        
        Args:
            content: Document content as string or bytes
            filename: Original filename
            content_type: MIME type or file extension
            
        Returns:
            Document object with extracted content
        """
        self.logger.info(f"Processing document: {filename}")
        
        try:
            # Extract text based on content type
            if content_type.lower() in ['.txt', '.md', 'text/plain']:
                text_content = self._process_text(content)
            elif content_type.lower() == '.json':
                text_content = self._process_json(content)
            elif content_type.lower() == '.html':
                text_content = self._process_html(content)
            elif content_type.lower() == '.pdf':
                text_content = self._process_pdf(content)
            elif content_type.lower() in ['.docx', '.doc']:
                text_content = self._process_docx(content)
            elif content_type.lower() in ['.xlsx', '.xls']:
                text_content = self._process_excel(content)
            elif content_type.lower() == '.csv':
                text_content = self._process_csv(content)
            else:
                # Default to treating as text
                text_content = self._process_text(content)
            
            # Create document object
            document = Document(
                filename=filename,
                content=text_content,
                content_type=content_type,
                file_size=len(content) if isinstance(content, (str, bytes)) else 0,
                processed_date=datetime.utcnow(),
                status="processing"
            )
            
            self.logger.info(f"Document processed successfully: {filename}")
            return document
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {e}")
            raise
    
    def _process_text(self, content: Union[str, bytes]) -> str:
        """Process plain text content"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        return self._clean_text(content)
    
    def _process_json(self, content: Union[str, bytes]) -> str:
        """Process JSON content by extracting text values"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        try:
            data = json.loads(content)
            text_parts = []
            
            def extract_text_from_json(obj):
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, str):
                            text_parts.append(f"{key}: {value}")
                        elif isinstance(value, (dict, list)):
                            extract_text_from_json(value)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text_from_json(item)
                elif isinstance(obj, str):
                    text_parts.append(obj)
            
            extract_text_from_json(data)
            return self._clean_text("\n".join(text_parts))
            
        except json.JSONDecodeError:
            return self._process_text(content)
    
    def _process_html(self, content: Union[str, bytes]) -> str:
        """Process HTML content by extracting text"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            text = soup.get_text(separator=' ', strip=True)
            return self._clean_text(text)
        except ImportError:
            # Fallback if BeautifulSoup is not available
            import re
            text = re.sub(r'<[^>]+>', ' ', content)
            return self._clean_text(text)
    
    def _process_pdf(self, content: Union[str, bytes]) -> str:
        """Process PDF content and extract text"""
        if isinstance(content, str):
            content = content.encode('utf-8')
        
        try:
            import pdfplumber
            import io
            
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = '\n'.join(text_parts)
                return self._clean_text(full_text)
                
        except ImportError:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = '\n'.join(text_parts)
                return self._clean_text(full_text)
                
            except Exception as e:
                self.logger.error(f"Error processing PDF: {e}")
                return "Error: Could not extract text from PDF"
                
        except Exception as e:
            self.logger.error(f"Error processing PDF with pdfplumber: {e}")
            return "Error: Could not extract text from PDF"
    
    def _process_docx(self, content: Union[str, bytes]) -> str:
        """Process DOCX/DOC content and extract text"""
        if isinstance(content, str):
            content = content.encode('utf-8')
        
        try:
            from docx import Document as DocxDocument
            import io
            
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            return "Error: Could not extract text from DOCX"
    
    def _process_excel(self, content: Union[str, bytes]) -> str:
        """Process Excel content and extract text"""
        if isinstance(content, str):
            content = content.encode('utf-8')
        
        try:
            import openpyxl
            import io
            
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(row_values).strip()
                    if row_text and row_text != " | " * (len(row_values) - 1):
                        text_parts.append(row_text)
            
            full_text = '\n'.join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            # Fallback to xlrd for older Excel formats
            try:
                import xlrd
                import io
                
                workbook = xlrd.open_workbook(file_contents=content)
                text_parts = []
                
                for sheet_idx in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    text_parts.append(f"Sheet: {sheet.name}")
                    
                    for row_idx in range(sheet.nrows):
                        row_values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                        row_text = " | ".join(row_values).strip()
                        if row_text:
                            text_parts.append(row_text)
                
                full_text = '\n'.join(text_parts)
                return self._clean_text(full_text)
                
            except Exception as e2:
                self.logger.error(f"Error processing Excel: {e}, {e2}")
                return "Error: Could not extract text from Excel file"
    
    def _process_csv(self, content: Union[str, bytes]) -> str:
        """Process CSV content and extract text"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        try:
            import csv
            import io
            
            csv_reader = csv.reader(io.StringIO(content))
            text_parts = []
            
            for row in csv_reader:
                row_text = " | ".join(str(cell).strip() for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
            
            full_text = '\n'.join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            self.logger.error(f"Error processing CSV: {e}")
            return self._process_text(content)  # Fallback to text processing
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters that might interfere
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\ufeff', '')  # Remove BOM
        
        # Remove very long lines that might be corrupted
        lines = text.split('\n')
        cleaned_lines = [line for line in lines if len(line) < 10000]
        
        return '\n'.join(cleaned_lines).strip()


class DocumentChunker:
    """Handles document chunking with various strategies"""
    
    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.logger = logger
    
    def chunk_document(self, document: Document, 
                      strategy: str = "semantic") -> List[DocumentChunk]:
        """
        Split document into chunks using specified strategy
        
        Args:
            document: Document to chunk
            strategy: Chunking strategy ('semantic', 'fixed', 'paragraph')
            
        Returns:
            List of DocumentChunk objects
        """
        self.logger.info(f"Chunking document {document.filename} with strategy: {strategy}")
        
        try:
            if strategy == "semantic":
                chunks = self._semantic_chunk(document)
            elif strategy == "fixed":
                chunks = self._fixed_size_chunk(document)
            elif strategy == "paragraph":
                chunks = self._paragraph_chunk(document)
            else:
                self.logger.warning(f"Unknown chunking strategy: {strategy}, using semantic")
                chunks = self._semantic_chunk(document)
            
            # Set chunk metadata
            for i, chunk in enumerate(chunks):
                chunk.source_document_id = document.id
                chunk.chunk_index = i
                chunk.metadata = {
                    'source_filename': document.filename,
                    'source_content_type': document.content_type,
                    'chunking_strategy': strategy,
                    'chunk_size': self.chunk_size,
                    'overlap': self.overlap
                }
            
            self.logger.info(f"Created {len(chunks)} chunks for document {document.filename}")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error chunking document {document.filename}: {e}")
            raise
    
    def _semantic_chunk(self, document: Document) -> List[DocumentChunk]:
        """Split text into semantically coherent chunks"""
        text = document.content
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence_len = len(sentence)
            
            # If adding this sentence would exceed chunk size, create a new chunk
            if current_size + sentence_len > self.chunk_size and current_chunk:
                chunks.append(DocumentChunk(content=current_chunk.strip()))
                
                # Handle overlap
                if self.overlap > 0:
                    current_chunk = current_chunk[-self.overlap:] + " " + sentence
                    current_size = len(current_chunk)
                else:
                    current_chunk = sentence
                    current_size = sentence_len
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_len
        
        # Add final chunk if there's content
        if current_chunk.strip():
            chunks.append(DocumentChunk(content=current_chunk.strip()))
        
        return chunks
    
    def _fixed_size_chunk(self, document: Document) -> List[DocumentChunk]:
        """Split text into fixed-size chunks with overlap"""
        text = document.content
        chunks = []
        
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk_text = text[start:end]
            
            if chunk_text.strip():
                chunks.append(DocumentChunk(content=chunk_text.strip()))
            
            start = end - self.overlap if self.overlap > 0 else end
        
        return chunks
    
    def _paragraph_chunk(self, document: Document) -> List[DocumentChunk]:
        """Split text into chunks by paragraphs"""
        text = document.content
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size, create a new chunk
            if len(current_chunk) + len(paragraph) > self.chunk_size and current_chunk:
                chunks.append(DocumentChunk(content=current_chunk.strip()))
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
        
        # Add final chunk if there's content
        if current_chunk.strip():
            chunks.append(DocumentChunk(content=current_chunk.strip()))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple heuristics"""
        import re
        
        # Simple sentence splitting using regex
        sentences = re.split(r'[.!?]+', text)
        
        # Filter out empty sentences and clean up
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences


class EnhancedDocumentChunker(DocumentChunker):
    """Enhanced chunker with full enterprise features"""
    
    def __init__(self, chunk_size: int = 512, overlap: int = 50):
        super().__init__(chunk_size, overlap)
        self.legal_validator = LegalChunkValidator()
        self.risk_detector = LegalRiskDetector()
        self.classifier = AdvancedDocumentClassifier()
        self.financial_chunker = FinancialDocumentChunker()
        
        # Document type detection patterns
        self.document_type_patterns = {
            'legal': [
                'agreement', 'contract', 'terms and conditions', 'license',
                'nda', 'non-disclosure', 'indemnif', 'liability', 'whereas',
                'party a', 'party b', 'counterpart', 'execution', 'witness'
            ],
            'financial': [
                'revenue', 'ebitda', 'balance sheet', 'income statement',
                'cash flow', 'quarterly', 'fiscal year', 'profit', 'loss',
                '$', 'usd', 'million', 'billion', 'percentage', '%'
            ],
            'code': [
                'function', 'class', 'import', 'def ', 'var ', 'const ',
                'return', 'if ', 'else', 'for ', 'while ', 'try', 'catch'
            ]
        }
    
    def detect_document_type(self, document: Document) -> str:
        """Detect document type for appropriate chunking strategy"""
        
        content_lower = document.content.lower()
        
        # Check file extension first
        if hasattr(document, 'filename'):
            filename = document.filename.lower()
            if any(ext in filename for ext in ['.py', '.js', '.java', '.cpp', '.rb']):
                return 'code'
            if any(ext in filename for ext in ['.xlsx', '.xls', '.csv']):
                return 'financial'
        
        # Content-based detection
        type_scores = {}
        
        for doc_type, patterns in self.document_type_patterns.items():
            score = sum(1 for pattern in patterns if pattern in content_lower)
            if score > 0:
                type_scores[doc_type] = score / len(patterns)
        
        if type_scores:
            detected_type = max(type_scores.items(), key=lambda x: x[1])[0]
            confidence = type_scores[detected_type]
            
            logger.info(f"Detected document type: {detected_type} (confidence: {confidence:.2f})")
            
            # Require minimum confidence for legal documents (safety)
            if detected_type == 'legal' and confidence < 0.1:
                return 'general'
                
            return detected_type
        
        return 'general'
    
    def chunk_document(self, document: Document, strategy: str = "auto") -> List[DocumentChunk]:
        """Enhanced chunking with full type awareness"""
        
        # Advanced document classification
        classification = self.classifier.classify_document(document)
        
        logger.info(
            f"Document classified as {classification.document_type.value} "
            f"(confidence: {classification.confidence:.2f})"
        )
        
        # Adjust chunking parameters based on classification
        original_chunk_size = self.chunk_size
        original_overlap = self.overlap
        
        self.chunk_size = classification.suggested_chunk_size
        self.overlap = classification.suggested_overlap
        
        try:
            # Apply type-specific chunking
            if classification.document_type == DocumentType.LEGAL:
                chunks = self._legal_safe_chunk(document)
            elif classification.document_type == DocumentType.FINANCIAL:
                # Use specialized financial chunker
                chunks = self.financial_chunker.chunk_financial_document(document)
                # Convert to DocumentChunk objects if needed
                chunks = self._convert_to_document_chunks(chunks, document)
            elif classification.document_type == DocumentType.CODE:
                chunks = self._code_aware_chunk(document)
            elif classification.document_type == DocumentType.TECHNICAL:
                chunks = self._technical_doc_chunk(document)
            elif classification.document_type == DocumentType.CONVERSATION:
                chunks = self._conversation_chunk(document)
            else:
                chunks = self._semantic_chunk(document)
            
            # Apply validation based on document type
            if classification.document_type == DocumentType.LEGAL:
                chunks = self._apply_legal_validation(chunks)
            elif classification.document_type == DocumentType.FINANCIAL:
                chunks = self._apply_financial_validation(chunks)
            
            # Add classification metadata to all chunks
            for chunk in chunks:
                if not hasattr(chunk, 'metadata'):
                    chunk.metadata = {}
                chunk.metadata.update({
                    'document_type': classification.document_type.value,
                    'classification_confidence': classification.confidence,
                    'detected_features': classification.detected_features,
                    'chunk_timestamp': datetime.utcnow().isoformat()
                })
            
            logger.info(f"Generated {len(chunks)} chunks for {classification.document_type.value} document")
            return chunks
            
        finally:
            # Restore original parameters
            self.chunk_size = original_chunk_size
            self.overlap = original_overlap
    
    def _legal_safe_chunk(self, document: Document) -> List[DocumentChunk]:
        """Legal-safe chunking that preserves clause integrity"""
        
        content = document.content
        chunks = []
        
        # First, identify high-risk areas
        risk_areas = []
        for i in range(0, len(content), 100):  # Scan in 100-char windows
            window = content[i:i+200]
            risk_assessment = self.risk_detector.assess_risk(window)
            
            if risk_assessment.level in [LegalRiskLevel.HIGH, LegalRiskLevel.CRITICAL]:
                risk_areas.append((i, i+200, risk_assessment))
        
        if not risk_areas:
            # No high-risk content, use standard semantic chunking
            return self._semantic_chunk(document)
        
        # For high-risk content, use enhanced section-based chunking
        sections = self._identify_legal_sections(content)
        
        if sections:
            # Analyze cross-references between sections
            section_groups = self._group_cross_referenced_sections(content, sections)
            
            # Chunk by section groups to preserve cross-reference integrity
            for group in section_groups:
                # Combine only the sections in the group (not the range between them)
                if len(group) == 1:
                    # Single section - use its exact boundaries
                    section = group[0]
                    chunk_content = content[section['start']:section['end']]
                else:
                    # Multiple sections - concatenate their content
                    group_content_parts = []
                    for section in sorted(group, key=lambda s: s['start']):
                        section_content = content[section['start']:section['end']]
                        group_content_parts.append(section_content.strip())
                    chunk_content = '\n\n'.join(group_content_parts)
                
                # Check if the combined content is too large for a single chunk
                if len(chunk_content) > self.chunk_size * 3:  # Allow larger chunks for legal integrity
                    # If too large, use smart splitting within the group
                    sub_chunks = self._smart_split_legal_group(chunk_content, group, content)
                    chunks.extend(sub_chunks)
                else:
                    # Ensure minimum viable size
                    if len(chunk_content.strip()) > 50:
                        chunk = DocumentChunk(
                            content=chunk_content.strip(),
                            source_document_id=document.id,
                            chunk_index=len(chunks)
                        )
                        
                        # Add legal-specific metadata
                        section_titles = [section.get('title', '') for section in group]
                        
                        # Calculate group boundaries for risk assessment
                        group_start = min(section['start'] for section in group)
                        group_end = max(section['end'] for section in group)
                        
                        chunk.metadata = {
                            'section_titles': section_titles,
                            'section_numbers': [section.get('number', '') for section in group],
                            'legal_safe': True,
                            'cross_referenced_group': True,
                            'contains_high_risk_terms': any(
                                area for area in risk_areas 
                                if area[0] >= group_start and area[1] <= group_end
                            )
                        }
                        
                        chunks.append(chunk)
        else:
            # Fallback: Use enhanced paragraph-based chunking with high-risk term protection
            chunks = self._high_risk_aware_paragraph_chunk(content, risk_areas, document)
        
        return chunks
    
    def _identify_legal_sections(self, content: str) -> List[Dict]:
        """Identify legal document sections and subsections"""
        import re
        
        sections = []
        
        # Common legal section patterns
        section_patterns = [
            r'^(\d+\.?\s+[A-Z][A-Z\s]+)$',          # "1. DEFINITIONS"
            r'^(Section\s+\d+\.?\d*\s*[:\-]?\s*[A-Z][^.]*)',  # "Section 1: Overview"
            r'^(Article\s+[IVXLC]+\s*[:\-]?\s*[A-Z][^.]*)',   # "Article IV: Termination"
            r'^(\([a-z]\)\s+[A-Z][^.]*)',                      # "(a) General Terms"
        ]
        
        lines = content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            for pattern in section_patterns:
                match = re.match(pattern, line_stripped, re.IGNORECASE)
                if match:
                    # End previous section
                    if current_section:
                        current_section['end'] = content.find(line)
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'title': match.group(1),
                        'number': self._extract_section_number(match.group(1)),
                        'start': content.find(line),
                        'end': len(content)  # Will be updated when next section found
                    }
                    break
        
        # Close last section
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_section_number(self, title: str) -> str:
        """Extract section number from title"""
        import re
        match = re.search(r'(\d+\.?\d*)', title)
        return match.group(1) if match else ""
    
    def _group_cross_referenced_sections(self, content: str, sections: List[Dict]) -> List[List[Dict]]:
        """Group sections that cross-reference each other"""
        import re
        
        # Find all cross-references in the content
        cross_ref_pattern = r'Section\s+(\d+\.?\d*)'
        cross_refs = list(re.finditer(cross_ref_pattern, content, re.IGNORECASE))
        
        # Create groups based on cross-references
        section_groups = []
        used_section_indices = set()
        
        for i, section in enumerate(sections):
            if i in used_section_indices:
                continue
                
            # Check if this section contains cross-references
            section_content = content[section['start']:section['end']]
            section_cross_refs = list(re.finditer(cross_ref_pattern, section_content, re.IGNORECASE))
            
            if section_cross_refs:
                # Create a group with this section and its references
                group = [section]
                used_section_indices.add(i)
                
                for cross_ref in section_cross_refs:
                    referenced_num = cross_ref.group(1)
                    
                    # Find the referenced section (handle subsection references)
                    for j, other_section in enumerate(sections):
                        if j not in used_section_indices:
                            # Exact match first
                            if other_section['number'] == referenced_num:
                                group.append(other_section)
                                used_section_indices.add(j)
                                break
                            # Handle subsection references (e.g., "3.2" matches section "3.")
                            elif '.' in referenced_num and other_section['number'].startswith(referenced_num.split('.')[0] + '.'):
                                group.append(other_section)
                                used_section_indices.add(j)
                                break
                
                section_groups.append(group)
            else:
                # Standalone section
                section_groups.append([section])
                used_section_indices.add(i)
        
        # Handle any remaining sections
        for i, section in enumerate(sections):
            if i not in used_section_indices:
                section_groups.append([section])
        
        return section_groups
    
    def _smart_split_legal_group(self, chunk_content: str, group: List[Dict], full_content: str) -> List[DocumentChunk]:
        """Smart splitting of large legal groups while preserving integrity"""
        
        chunks = []
        
        # If the group is too large, split at natural boundaries but avoid high-risk terms
        high_risk_terms = ['liability', 'indemnif', 'except', 'provided', 'subject to']
        
        # Split by paragraphs first
        paragraphs = chunk_content.split('\n\n')
        current_chunk = ""
        current_size = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            # Check if adding this paragraph would exceed size and we have content
            if current_size + para_size > self.chunk_size and current_chunk:
                # Before splitting, check if current chunk ends with high-risk terms
                if self._ends_with_high_risk_pattern(current_chunk):
                    # Don't split here, continue building the chunk
                    current_chunk += "\n\n" + para
                    current_size += para_size
                else:
                    # Safe to split here
                    if current_chunk.strip():
                        chunk = DocumentChunk(
                            content=current_chunk.strip(),
                            source_document_id="",  # Will be set by caller
                            chunk_index=len(chunks)
                        )
                        chunk.metadata = {
                            'legal_safe_split': True,
                            'from_large_group': True
                        }
                        chunks.append(chunk)
                    
                    current_chunk = para
                    current_size = para_size
            else:
                current_chunk += "\n\n" + para if current_chunk else para
                current_size += para_size
        
        # Add the final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                source_document_id="",  # Will be set by caller
                chunk_index=len(chunks)
            )
            chunk.metadata = {
                'legal_safe_split': True,
                'from_large_group': True
            }
            chunks.append(chunk)
        
        return chunks
    
    def _ends_with_high_risk_pattern(self, content: str) -> bool:
        """Check if content ends with patterns that shouldn't be split"""
        import re
        
        content_end = content.strip().lower()
        
        # Patterns that indicate incomplete legal clauses
        risky_endings = [
            r'except\s*$',
            r'provided\s*$',
            r'subject\s+to\s*$',
            r'section\s+\d+\.?\d*\s*$',
            r'which\s+(limits?|provides?|states?)\s*$',
            r'liability\s+to\s*$',
            r'as\s+provided\s*$'
        ]
        
        return any(re.search(pattern, content_end) for pattern in risky_endings)
    
    def _high_risk_aware_paragraph_chunk(self, content: str, risk_areas: List, document: Document) -> List[DocumentChunk]:
        """Paragraph-based chunking that's aware of high-risk terms"""
        
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if adding this paragraph would exceed chunk size
            if len(current_chunk) + len(para) > self.chunk_size * 2 and current_chunk:  # Larger chunks for legal
                # Before splitting, check for high-risk patterns
                if self._would_split_high_risk_term(current_chunk, para):
                    # Don't split, continue building chunk
                    current_chunk += "\n\n" + para
                else:
                    # Safe to split
                    if current_chunk.strip():
                        chunk = DocumentChunk(
                            content=current_chunk.strip(),
                            source_document_id=document.id,
                            chunk_index=len(chunks)
                        )
                        chunk.metadata = {
                            'legal_safe': True,
                            'high_risk_aware_split': True
                        }
                        chunks.append(chunk)
                    current_chunk = para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Add final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                source_document_id=document.id,
                chunk_index=len(chunks)
            )
            chunk.metadata = {
                'legal_safe': True,
                'high_risk_aware_split': True
            }
            chunks.append(chunk)
        
        return chunks
    
    def _would_split_high_risk_term(self, current_chunk: str, next_para: str) -> bool:
        """Check if splitting here would separate high-risk terms from their context"""
        
        # Check if current chunk ends with a reference to the next paragraph
        current_end = current_chunk.strip().lower()
        next_start = next_para.strip().lower()
        
        # Look for cross-references
        import re
        if re.search(r'section\s+\d+\.?\d*', current_end):
            # Check if next paragraph is the referenced section
            if re.match(r'\d+\.?\d*\s+\w+', next_start):
                return True
        
        # Check for liability-related continuations
        liability_patterns = [
            (r'liability', r'(limitations?|caps?|exclusions?)'),
            (r'indemnif', r'(provisions?|terms?|obligations?)'),
            (r'except\s*$', r'(as|for|when)'),
        ]
        
        for end_pattern, start_pattern in liability_patterns:
            if re.search(end_pattern, current_end) and re.search(start_pattern, next_start):
                return True
        
        return False
    
    def _convert_to_document_chunks(self, chunks: List[Dict], document) -> List[DocumentChunk]:
        """Convert dictionary chunks to DocumentChunk objects"""
        
        document_chunks = []
        for chunk_dict in chunks:
            chunk = DocumentChunk(
                content=chunk_dict['content'],
                source_document_id=chunk_dict['source_document_id'],
                chunk_index=chunk_dict['chunk_index']
            )
            chunk.metadata = chunk_dict.get('metadata', {})
            document_chunks.append(chunk)
        
        return document_chunks
    
    def _apply_legal_validation(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Apply legal validation to chunks"""
        
        validated_chunks = []
        
        for chunk in chunks:
            validation_result = self.legal_validator.validate_legal_chunk(
                chunk.content, 
                f"{chunk.source_document_id}_{chunk.chunk_index}"
            )
            
            if not validation_result["is_valid"]:
                logger.error(f"Legal validation failed for chunk {chunk.chunk_index}: {validation_result['errors']}")
                # Could implement auto-repair here or flag for manual review
                continue
            
            # Add legal metadata
            chunk.metadata = chunk.metadata or {}
            chunk.metadata.update(validation_result["metadata_additions"])
            
            if validation_result["warnings"]:
                logger.warning(f"Legal warnings for chunk {chunk.chunk_index}: {validation_result['warnings']}")
            
            validated_chunks.append(chunk)
        
        return validated_chunks
    
    def _apply_financial_validation(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Apply financial-specific validation"""
        import re
        
        validated_chunks = []
        
        for chunk in chunks:
            # Check for incomplete calculations
            if self._has_incomplete_calculation(chunk.content):
                logger.warning(f"Potential incomplete calculation in chunk {chunk.chunk_index}")
                chunk.metadata = chunk.metadata or {}
                chunk.metadata['calculation_warning'] = True
            
            # Check for orphaned numbers
            if self._has_orphaned_numbers(chunk.content):
                logger.warning(f"Orphaned numbers detected in chunk {chunk.chunk_index}")
                chunk.metadata = chunk.metadata or {}
                chunk.metadata['orphaned_numbers_warning'] = True
            
            validated_chunks.append(chunk)
        
        return validated_chunks
    
    def _has_incomplete_calculation(self, content: str) -> bool:
        """Check for incomplete calculations"""
        import re
        
        # Look for dangling operators
        if re.search(r'[=+\-*/]\s*$', content.strip()):
            return True
        
        # Look for references to missing variables
        if re.search(r'\b(see\s+above|as\s+calculated|from\s+table)\s*$', content.strip(), re.IGNORECASE):
            return True
        
        return False
    
    def _has_orphaned_numbers(self, content: str) -> bool:
        """Check for numbers without context"""
        import re
        
        # Find standalone currency amounts
        amounts = re.findall(r'\$[\d,]+\.?\d*', content)
        
        # Simple heuristic: if more than 3 amounts without context words
        if len(amounts) > 3:
            context_words = ['revenue', 'cost', 'profit', 'total', 'budget', 'expense']
            has_context = any(word in content.lower() for word in context_words)
            if not has_context:
                return True
        
        return False
    
    def _technical_doc_chunk(self, document: Document) -> List[DocumentChunk]:
        """Technical documentation chunking"""
        import re
        
        content = document.content
        
        # Look for headers (markdown style)
        headers = list(re.finditer(r'^#{1,6}\s+(.+)$', content, re.MULTILINE))
        
        if len(headers) > 1:
            # Chunk by sections
            chunks = []
            for i, header in enumerate(headers):
                start = header.start()
                end = headers[i + 1].start() if i + 1 < len(headers) else len(content)
                
                section_content = content[start:end].strip()
                if len(section_content) > 100:
                    chunk = DocumentChunk(
                        content=section_content,
                        source_document_id=document.id,
                        chunk_index=len(chunks)
                    )
                    chunk.metadata = {
                        'section_header': header.group(1),
                        'header_level': len(header.group().split()[0]),  # Count #'s
                        'contains_code': '```' in section_content
                    }
                    chunks.append(chunk)
            
            return chunks
        else:
            # Use paragraph-based chunking
            return self._semantic_chunk(document)
    
    def _conversation_chunk(self, document: Document) -> List[DocumentChunk]:
        """Conversation/transcript chunking"""
        import re
        
        content = document.content
        
        # Look for speaker patterns
        speaker_pattern = r'^([A-Z][a-z]+)\s*:\s*(.+)$'
        turns = []
        
        for match in re.finditer(speaker_pattern, content, re.MULTILINE):
            turns.append({
                'speaker': match.group(1),
                'content': match.group(2),
                'position': match.start()
            })
        
        if len(turns) > 3:
            # Group turns into conversation segments
            chunks = []
            current_segment = []
            current_length = 0
            
            for turn in turns:
                turn_length = len(turn['content'])
                
                if current_length + turn_length > self.chunk_size and current_segment:
                    # Create chunk from current segment
                    chunk_content = self._format_conversation_segment(current_segment)
                    chunk = DocumentChunk(
                        content=chunk_content,
                        source_document_id=document.id,
                        chunk_index=len(chunks)
                    )
                    chunk.metadata = {
                        'conversation_turns': len(current_segment),
                        'speakers': list(set(t['speaker'] for t in current_segment))
                    }
                    chunks.append(chunk)
                    
                    # Start new segment with overlap
                    current_segment = current_segment[-1:] if current_segment else []
                    current_length = len(current_segment[0]['content']) if current_segment else 0
                
                current_segment.append(turn)
                current_length += turn_length
            
            # Don't forget the last segment
            if current_segment:
                chunk_content = self._format_conversation_segment(current_segment)
                chunk = DocumentChunk(
                    content=chunk_content,
                    source_document_id=document.id,
                    chunk_index=len(chunks)
                )
                chunk.metadata = {
                    'conversation_turns': len(current_segment),
                    'speakers': list(set(t['speaker'] for t in current_segment))
                }
                chunks.append(chunk)
            
            return chunks
        else:
            # No clear conversation structure
            return self._semantic_chunk(document)
    
    def _format_conversation_segment(self, segment: List[Dict]) -> str:
        """Format conversation segment for chunk"""
        
        formatted_lines = []
        for turn in segment:
            formatted_lines.append(f"{turn['speaker']}: {turn['content']}")
        
        return '\n'.join(formatted_lines)
    
    def _code_aware_chunk(self, document: Document) -> List[DocumentChunk]:
        """Code-aware chunking (basic implementation)"""
        
        # For code files, use smaller chunks focused on functions/classes
        original_chunk_size = self.chunk_size
        self.chunk_size = 400  # Smaller chunks for code
        
        chunks = self._semantic_chunk(document)
        
        # Restore original chunk size
        self.chunk_size = original_chunk_size
        
        return chunks


class DocumentStore:
    """Handles document storage and retrieval"""
    
    def __init__(self, db_path: str = "data/rag_documents.db", 
                 documents_dir: str = "data/documents"):
        self.db_path = db_path
        self.documents_dir = Path(documents_dir)
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger
        self._init_database()
    
    def _init_database(self):
        """Initialize SQLite database for document metadata"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS documents (
                        id TEXT PRIMARY KEY,
                        filename TEXT NOT NULL,
                        content_type TEXT,
                        file_size INTEGER,
                        upload_date TEXT,
                        processed_date TEXT,
                        metadata TEXT,
                        chunk_count INTEGER DEFAULT 0,
                        status TEXT DEFAULT 'pending'
                    )
                """)
                
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS document_chunks (
                        chunk_id TEXT PRIMARY KEY,
                        document_id TEXT,
                        chunk_index INTEGER,
                        content TEXT,
                        metadata TEXT,
                        relevance_score REAL DEFAULT 0.0,
                        created_at TEXT,
                        FOREIGN KEY (document_id) REFERENCES documents (id)
                    )
                """)
                
                # Create indexes for performance
                conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chunks_document_id 
                    ON document_chunks(document_id)
                """)
                
                conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chunks_relevance 
                    ON document_chunks(relevance_score DESC)
                """)
                
                conn.commit()
                
        except Exception as e:
            self.logger.error(f"Error initializing database: {e}")
            raise
    
    def store_document(self, document: Document, 
                      chunks: List[DocumentChunk]) -> bool:
        """Store document and its chunks"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # Store document metadata
                conn.execute("""
                    INSERT OR REPLACE INTO documents 
                    (id, filename, content_type, file_size, upload_date, 
                     processed_date, metadata, chunk_count, status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    document.id,
                    document.filename,
                    document.content_type,
                    document.file_size,
                    document.upload_date.isoformat(),
                    document.processed_date.isoformat() if document.processed_date else None,
                    json.dumps(document.metadata),
                    len(chunks),
                    document.status
                ))
                
                # Store chunks
                for chunk in chunks:
                    conn.execute("""
                        INSERT OR REPLACE INTO document_chunks 
                        (chunk_id, document_id, chunk_index, content, metadata, created_at)
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, (
                        chunk.chunk_id,
                        chunk.source_document_id,
                        chunk.chunk_index,
                        chunk.content,
                        json.dumps(chunk.metadata),
                        chunk.created_at.isoformat()
                    ))
                
                # Store full document content separately
                doc_file_path = self.documents_dir / f"{document.id}.json"
                with open(doc_file_path, 'w', encoding='utf-8') as f:
                    json.dump({
                        'document': document.to_dict(),
                        'content': document.content,
                        'chunks': [chunk.to_dict() for chunk in chunks]
                    }, f, ensure_ascii=False, indent=2)
                
                conn.commit()
                
            self.logger.info(f"Stored document {document.id} with {len(chunks)} chunks")
            return True
            
        except Exception as e:
            self.logger.error(f"Error storing document: {e}")
            return False
    
    def retrieve_document(self, document_id: str) -> Optional[Document]:
        """Retrieve document by ID"""
        try:
            doc_file_path = self.documents_dir / f"{document_id}.json"
            if not doc_file_path.exists():
                return None
            
            with open(doc_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            doc_data = data['document']
            document = Document(
                id=doc_data['id'],
                filename=doc_data['filename'],
                content=data['content'],
                content_type=doc_data['content_type'],
                file_size=doc_data['file_size'],
                metadata=doc_data['metadata'],
                upload_date=datetime.fromisoformat(doc_data['upload_date']),
                processed_date=datetime.fromisoformat(doc_data['processed_date']) if doc_data['processed_date'] else None,
                status=doc_data['status']
            )
            
            # Load chunks
            for chunk_data in data['chunks']:
                chunk = DocumentChunk.from_dict(chunk_data)
                document.chunks.append(chunk)
            
            return document
            
        except Exception as e:
            self.logger.error(f"Error retrieving document {document_id}: {e}")
            return None
    
    def get_chunks_by_document_id(self, document_id: str) -> List[DocumentChunk]:
        """Get all chunks for a document"""
        chunks = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    SELECT chunk_id, chunk_index, content, metadata, created_at
                    FROM document_chunks 
                    WHERE document_id = ?
                    ORDER BY chunk_index
                """, (document_id,))
                
                for row in cursor.fetchall():
                    chunk = DocumentChunk(
                        chunk_id=row[0],
                        source_document_id=document_id,
                        chunk_index=row[1],
                        content=row[2],
                        metadata=json.loads(row[3]),
                        created_at=datetime.fromisoformat(row[4])
                    )
                    chunks.append(chunk)
                    
        except Exception as e:
            self.logger.error(f"Error retrieving chunks for document {document_id}: {e}")
        
        return chunks
    
    def search_documents(self, query: str, limit: int = 10) -> List[Dict]:
        """Simple text search across documents"""
        results = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    SELECT id, filename, content_type, upload_date, chunk_count, status
                    FROM documents 
                    WHERE filename LIKE ? OR metadata LIKE ?
                    ORDER BY upload_date DESC
                    LIMIT ?
                """, (f"%{query}%", f"%{query}%", limit))
                
                for row in cursor.fetchall():
                    results.append({
                        'id': row[0],
                        'filename': row[1],
                        'content_type': row[2],
                        'upload_date': row[3],
                        'chunk_count': row[4],
                        'status': row[5]
                    })
                    
        except Exception as e:
            self.logger.error(f"Error searching documents: {e}")
        
        return results
    
    def list_documents(self, limit: int = 50, offset: int = 0) -> List[Dict]:
        """List all documents"""
        results = []
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    SELECT id, filename, content_type, upload_date, chunk_count, status
                    FROM documents 
                    ORDER BY upload_date DESC
                    LIMIT ? OFFSET ?
                """, (limit, offset))
                
                for row in cursor.fetchall():
                    results.append({
                        'id': row[0],
                        'filename': row[1],
                        'content_type': row[2],
                        'upload_date': row[3],
                        'chunk_count': row[4],
                        'status': row[5]
                    })
                    
        except Exception as e:
            self.logger.error(f"Error listing documents: {e}")
        
        return results
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                # Delete chunks first
                conn.execute("DELETE FROM document_chunks WHERE document_id = ?", (document_id,))
                
                # Delete document
                conn.execute("DELETE FROM documents WHERE id = ?", (document_id,))
                
                # Delete document file
                doc_file_path = self.documents_dir / f"{document_id}.json"
                if doc_file_path.exists():
                    doc_file_path.unlink()
                
                conn.commit()
                
            self.logger.info(f"Deleted document {document_id}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error deleting document {document_id}: {e}")
            return False
